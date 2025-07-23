"""
Content extractors for different file types
"""
import os
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Optional, Dict, Any
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
from pathlib import Path

logger = logging.getLogger(__name__)

@dataclass
class ExtractedContent:
    """Container for extracted content and metadata"""
    text: str
    page_count: int = 0
    confidence: float = 1.0
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}

class BaseExtractor(ABC):
    """Base class for content extractors"""
    
    @abstractmethod
    def extract(self, file_path: str) -> ExtractedContent:
        """Extract content from file"""
        pass
    
    @abstractmethod
    def supports_format(self, file_path: str) -> bool:
        """Check if this extractor supports the file format"""
        pass
    
    def validate_file(self, file_path: str) -> bool:
        """Validate that file exists and is readable"""
        return os.path.exists(file_path) and os.path.isfile(file_path)

class PDFExtractor(BaseExtractor):
    """Extract text from PDF files using multiple methods"""
    
    def extract(self, file_path: str) -> ExtractedContent:
        try:
            if not self.validate_file(file_path):
                raise ValueError(f"Invalid file: {file_path}")

            # Try PyMuPDF first (fastest)
            text, page_count = self._extract_with_pymupdf(file_path)
            confidence = 0.95

            # Fallback to pdfplumber if PyMuPDF fails
            if not text.strip():
                text, page_count = self._extract_with_pdfplumber(file_path)
                confidence = 0.9

            # Fallback to PyPDF2 if both fail
            if not text.strip():
                text, page_count = self._extract_with_pypdf2(file_path)
                confidence = 0.7

            return ExtractedContent(
                text=text,
                page_count=page_count,
                confidence=confidence,
                metadata={'extraction_method': 'pdf', 'file_size': os.path.getsize(file_path)}
            )

        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            return ExtractedContent(text="", confidence=0.0)
    
    def _extract_with_pdfplumber(self, file_path: str) -> tuple:
        """Extract using pdfplumber (better for tables)"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return '\n'.join(text_parts), len(pdf.pages)
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
            return "", 0
    
    def _extract_with_pypdf2(self, file_path: str) -> tuple:
        """Extract using PyPDF2 (fallback method)"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return '\n'.join(text_parts), len(reader.pages)
                
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {e}")
            return "", 0
    
    def _extract_with_pymupdf(self, file_path: str) -> tuple:
        """Extract using PyMuPDF (fitz) - very fast"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text_parts = []
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
            return '\n'.join(text_parts), len(doc)
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {file_path}: {e}")
            return "", 0
    
    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith('.pdf')

class WordExtractor(BaseExtractor):
    """Extract text from Word documents (.docx)"""
    
    def extract(self, file_path: str) -> ExtractedContent:
        try:
            if not self.validate_file(file_path):
                raise ValueError(f"Invalid file: {file_path}")
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            page_count = len(doc.sections) if doc.sections else 1
            
            return ExtractedContent(
                text=full_text,
                page_count=page_count,
                confidence=0.95,
                metadata={'extraction_method': 'docx', 'file_size': os.path.getsize(file_path)}
            )
            
        except Exception as e:
            logger.error(f"Error extracting Word document {file_path}: {e}")
            return ExtractedContent(text="", confidence=0.0)
    
    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))

class TextExtractor(BaseExtractor):
    """Extract text from plain text files"""
    
    def extract(self, file_path: str) -> ExtractedContent:
        try:
            if not self.validate_file(file_path):
                raise ValueError(f"Invalid file: {file_path}")
            
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                raise ValueError("Could not decode text file")
            
            # Estimate page count (assuming ~500 words per page)
            word_count = len(text.split())
            page_count = max(1, word_count // 500)
            
            return ExtractedContent(
                text=text,
                page_count=page_count,
                confidence=1.0,
                metadata={'extraction_method': 'text', 'file_size': os.path.getsize(file_path)}
            )
            
        except Exception as e:
            logger.error(f"Error extracting text file {file_path}: {e}")
            return ExtractedContent(text="", confidence=0.0)
    
    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.txt', '.md', '.csv'))

class ExcelExtractor(BaseExtractor):
    """Extract text from Excel files"""
    
    def extract(self, file_path: str) -> ExtractedContent:
        try:
            if not self.validate_file(file_path):
                raise ValueError(f"Invalid file: {file_path}")
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert dataframe to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False, na_rep='')
                text_parts.append(sheet_text)
            
            full_text = '\n\n'.join(text_parts)
            page_count = len(excel_file.sheet_names)
            
            return ExtractedContent(
                text=full_text,
                page_count=page_count,
                confidence=0.8,
                metadata={'extraction_method': 'excel', 'file_size': os.path.getsize(file_path)}
            )
            
        except Exception as e:
            logger.error(f"Error extracting Excel file {file_path}: {e}")
            return ExtractedContent(text="", confidence=0.0)
    
    def supports_format(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.xlsx', '.xls'))

class ContentExtractorFactory:
    """Factory for creating appropriate content extractors"""
    
    _extractors = [
        PDFExtractor(),
        WordExtractor(),
        TextExtractor(),
        ExcelExtractor()
    ]
    
    @classmethod
    def get_extractor(cls, file_path: str) -> Optional[BaseExtractor]:
        """Get appropriate extractor for file format"""
        for extractor in cls._extractors:
            if extractor.supports_format(file_path):
                return extractor
        return None
    
    @classmethod
    def get_supported_formats(cls) -> list:
        """Get list of supported file formats"""
        formats = ['.pdf', '.docx', '.doc', '.txt', '.md', '.csv', '.xlsx', '.xls']
        return formats

# Helper function for external use
def extract_content(file_path: str) -> ExtractedContent:
    """Extract content from any supported file format"""
    extractor = ContentExtractorFactory.get_extractor(file_path)
    if extractor:
        return extractor.extract(file_path)
    else:
        logger.error(f"No extractor available for file: {file_path}")
        return ExtractedContent(text="", confidence=0.0, metadata={'error': 'Unsupported file format'})